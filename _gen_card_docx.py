# -*- coding: utf-8 -*-
"""Generate backend/mdFile/新卡牌专属效果.docx from the markdown tables."""
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path

ROOT = Path(__file__).resolve().parent
MD = ROOT / "backend" / "mdFile" / "新卡牌专属效果.md"
OUT = ROOT / "backend" / "mdFile" / "新卡牌专属效果.docx"
V2_MD = ROOT / "backend" / "mdFile" / "新卡牌专属效果（第二版）.md"
V2_OUT = ROOT / "backend" / "mdFile" / "新卡牌专属效果（第二版）.docx"
HEADERS = ["编号", "名称", "费用", "类型", "专属效果", "设计说明"]
CENTER_COLS = {0, 2, 3}


def set_run_font(run, size=10.5, bold=False, color=None, name="宋体"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:eastAsia"), name)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "8FAADC")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, v in (("top", 40), ("left", 60), ("bottom", 40), ("right", 60)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_valign(cell):
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    cell._tc.get_or_add_tcPr().append(vAlign)


def write_cell(cell, text, *, header=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (header or center) else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    if header:
        set_run_font(run, 10.5, True, RGBColor(0xC0, 0x00, 0x00), "微软雅黑")
        run.italic = True
        shade_cell(cell, "E2EFDA")
    else:
        set_run_font(run, 10.5, False, RGBColor(0x2D, 0x2D, 0x2D), "宋体")
    set_cell_border(cell)
    set_cell_margins(cell)
    set_cell_valign(cell)


def add_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    widths = [Cm(2.2), Cm(2.6), Cm(1.5), Cm(1.8), Cm(8.6), Cm(7.3)]
    for i, h in enumerate(HEADERS):
        write_cell(table.rows[0].cells[i], h, header=True, center=True)
        table.rows[0].cells[i].width = widths[i]
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            write_cell(cell, val, center=(c in CENTER_COLS))
            cell.width = widths[c]
            if r % 2 == 0:
                shade_cell(cell, "F7FAFC")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 16, True, RGBColor(0x1F, 0x4E, 0x79), "微软雅黑")
    else:
        set_run_font(run, 13, True, RGBColor(0x2E, 0x75, 0xB6), "微软雅黑")


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, 11, False, RGBColor(0x33, 0x33, 0x33), "宋体")


def parse_md_tables(text):
    sections = []
    heading = None
    note = None
    rows = []
    in_table = False
    for line in text.splitlines():
        if line.startswith("## "):
            if heading and rows:
                sections.append((heading, note, rows))
            heading = line[3:].replace("`", "").strip()
            note = None
            rows = []
            in_table = False
            continue
        if line.startswith("|") and "编号" in line and "名称" in line:
            in_table = True
            continue
        if in_table and line.startswith("|") and set(line.replace("|", "").strip()) <= set("-: "):
            continue
        if in_table and line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if len(cells) >= 6:
                rows.append(cells[:6])
            continue
        if in_table and not line.startswith("|"):
            in_table = False
        if heading and not rows and line.strip() and not line.startswith("|") and not line.startswith("#") and not line.startswith("---"):
            note = line.strip()
    if heading and rows:
        sections.append((heading, note, rows))
    return sections


def main():
    sections = parse_md_tables(MD.read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("新卡牌专属效果")
    set_run_font(run, 22, True, RGBColor(0x1F, 0x4E, 0x79), "微软雅黑")

    add_body(doc, "记录需胜利解锁的收藏卡（及图鉴样例卡 T-01）当前专属效果。费用、部门、卡牌类型不变。效果以结构化配置为准，与库内 cards.description 一致。")
    add_body(doc, "范围：O-16～O-19、S-07～S-24、P-05～P-20、O-20～O-44、T-01、H-01～H-03、B-01～B-07")
    add_body(doc, "不含：初始 27 张旧卡（S-01～S-06、P-01～P-04、O-01～O-15、L-01 / L-02）")

    add_heading(doc, "设计原则", 2)
    for line in [
        "整套机制（类型 + 时机 + 延迟 + 数值 + 目标）互不重复，也不复刻旧卡整套。",
        "公共卡强度按费用：0 费≈2、1 费≈3、2 费≈4、3 费≈6。",
        "老板卡同费明显高于公共卡（约高一档费用）。",
        "销售偏输出，采购偏防御，公共可混搭；目标「自己 / 一名玩家 / 双方」视为不同效果。",
    ]:
        add_body(doc, "· " + line)

    for heading, note, rows in sections:
        if heading == "修订记录":
            continue
        add_heading(doc, heading, 2)
        if note:
            add_body(doc, note)
        add_table(doc, rows)

    rec_heading = None
    rec_rows = []
    in_rec = False
    for line in MD.read_text(encoding="utf-8").splitlines():
        if line.startswith("## ") and "修订记录" in line:
            rec_heading = "修订记录"
            continue
        if rec_heading and line.startswith("|") and "日期" in line:
            in_rec = True
            continue
        if in_rec and line.startswith("|") and set(line.replace("|", "").strip()) <= set("-: "):
            continue
        if in_rec and line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rec_rows.append(cells)
        elif in_rec:
            break
    if rec_heading:
        add_heading(doc, rec_heading, 2)
        rec = doc.add_table(rows=1 + len(rec_rows), cols=2)
        rec.alignment = WD_TABLE_ALIGNMENT.CENTER
        write_cell(rec.rows[0].cells[0], "日期", header=True, center=True)
        write_cell(rec.rows[0].cells[1], "说明", header=True, center=True)
        for i, row in enumerate(rec_rows, start=1):
            write_cell(rec.rows[i].cells[0], row[0], center=True)
            write_cell(rec.rows[i].cells[1], row[1] if len(row) > 1 else "")

    tmp = OUT.with_suffix(".tmp.docx")
    doc.save(tmp)
    try:
        tmp.replace(OUT)
        print("wrote", OUT, "sections", len(sections))
    except PermissionError:
        alt = OUT.with_name("新卡牌专属效果-新表.docx")
        tmp.replace(alt)
        print("locked, wrote", alt, "sections", len(sections))


def save_doc(doc, out: Path):
    tmp = out.with_suffix(".tmp.docx")
    doc.save(tmp)
    try:
        tmp.replace(out)
        print("wrote", out)
    except PermissionError:
        alt = out.with_name(out.stem + "-新表.docx")
        tmp.replace(alt)
        print("locked, wrote", alt)


def main_v2():
    sections = parse_md_tables(V2_MD.read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("新卡牌专属效果（第二版）")
    set_run_font(run, 22, True, RGBColor(0x1F, 0x4E, 0x79), "微软雅黑")

    add_body(doc, "记录自 2026-08-24 起新增的收藏卡专属效果。第一版（此前全部收藏卡）见同目录《新卡牌专属效果.docx》。")
    add_body(doc, "费用、部门、卡牌类型以入库配置为准；效果与库内 cards.description 一致。表格列：编号、名称、费用、类型、专属效果、设计说明。")

    add_heading(doc, "设计原则", 2)
    for line in [
        "只收录今日起新加的卡，不回写第一版已有卡面。",
        "整套机制（类型 + 时机 + 延迟 + 数值 + 目标）不与旧卡、第一版新卡重复。",
        "老板卡同费明显高于公共卡（约高一档费用）。",
    ]:
        add_body(doc, "· " + line)

    for heading, note, rows in sections:
        if heading == "修订记录":
            continue
        add_heading(doc, heading, 2)
        if note:
            add_body(doc, note)
        add_table(doc, rows)

    text = V2_MD.read_text(encoding="utf-8")
    rec_heading = None
    rec_rows = []
    in_rec = False
    for line in text.splitlines():
        if line.startswith("## ") and "修订记录" in line:
            rec_heading = "修订记录"
            continue
        if rec_heading and line.startswith("|") and "日期" in line:
            in_rec = True
            continue
        if in_rec and line.startswith("|") and set(line.replace("|", "").strip()) <= set("-: "):
            continue
        if in_rec and line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rec_rows.append(cells)
        elif in_rec:
            break
    if rec_heading:
        add_heading(doc, rec_heading, 2)
        rec = doc.add_table(rows=1 + len(rec_rows), cols=2)
        rec.alignment = WD_TABLE_ALIGNMENT.CENTER
        write_cell(rec.rows[0].cells[0], "日期", header=True, center=True)
        write_cell(rec.rows[0].cells[1], "说明", header=True, center=True)
        for i, row in enumerate(rec_rows, start=1):
            write_cell(rec.rows[i].cells[0], row[0], center=True)
            write_cell(rec.rows[i].cells[1], row[1] if len(row) > 1 else "")

    save_doc(doc, V2_OUT)


CURRENT_OUT = ROOT / "backend" / "mdFile" / "卡牌顾客与霸凌者特效.docx"


def add_flex_table(doc, headers, rows, widths, center_cols=None):
    center_cols = center_cols or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, header=True, center=True)
        table.rows[0].cells[i].width = widths[i]
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            write_cell(cell, str(val), center=(c in center_cols))
            cell.width = widths[c]
            if r % 2 == 0:
                shade_cell(cell, "F7FAFC")
    return table


def main_current():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("卡牌、顾客效果与霸凌者特效")
    set_run_font(run, 22, True, RGBColor(0x1F, 0x4E, 0x79), "微软雅黑")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("当前对局配置 · 2026-09-01")
    set_run_font(run, 12, False, RGBColor(0x5B, 0x5B, 0x5B), "微软雅黑")

    add_body(doc, "本文记录当前对局正在使用的内容：销售/采购成员卡、四名顾客、以及与顾客一对一绑定的霸凌者特效。公共部与中立卡已下架，不再进入组牌。")
    add_body(doc, "更早的收藏卡专属效果仍见同目录《新卡牌专属效果.docx》《新卡牌专属效果（第二版）.docx》。")

    add_heading(doc, "对局与组牌", 1)
    for line in [
        "目标：把霸凌者血量打到 0。双方都倒下则失败。销售初始 50 血，采购初始 75 血。",
        "霸凌者开局血量 100～110 随机，基础攻击 5。打霸凌者时先扣其本回合护盾，再扣血。",
        "组牌只用本部门成员卡：本部门基础卡每种 2 张 + 已解锁收藏卡各 1 张。不含对方部门、公共部、中立卡。",
        "满解锁时销售约 15 张，采购约 18 张；解锁不足可以少于 30 张。",
        "每回合 3 点调用机会、抽满 5 张手牌。双方结束后霸凌者出手。玩家护盾结算后清空。",
        "顾客效果从第 2 回合开始掷骰；霸凌者特效从第 1 回合就生效。两套概率分开掷，不会绑成同时触发。",
        "同名角色的采购卡与销售卡是不同卡。",
    ]:
        add_body(doc, "· " + line)

    card_headers = ["编号", "名称", "费用", "获取", "类型", "专属效果", "设计说明"]
    card_widths = [Cm(2.0), Cm(2.6), Cm(1.4), Cm(1.6), Cm(1.6), Cm(7.8), Cm(7.0)]
    card_center = {0, 2, 3, 4}

    add_heading(doc, "销售成员卡 S-35～S-43", 1)
    add_body(doc, "轻铠剑士立绘。纯攻不持盾；带盾效果才持盾。兵器按攻击/护盾强弱缩放，每人不同。基础卡每局带 2 张，收藏卡需胜利解锁。")
    add_flex_table(
        doc,
        card_headers,
        [
            ["S-35", "Colin", "0", "基础", "攻击", "立即造成 2 点伤害", "短单手臂剑（最小）。"],
            ["S-36", "Carl", "1", "基础", "攻击", "立即造成 4 点伤害", "中等镰刀剑 falchion。"],
            ["S-37", "Nico", "2", "基础", "攻击", "立即造成 9 点伤害", "双手丹麦大斧。"],
            ["S-38", "Amy", "3", "收藏", "攻击", "立即造成 15 点伤害", "很大的双手大剑。"],
            ["S-43", "Kinyond", "3", "收藏", "攻击", "50% 概率立即造成 30 点伤害", "最大长柄斧 pollaxe。未中则 0 伤。"],
            ["S-39", "Kade", "0", "收藏", "攻击", "立即造成 2 点伤害，并为自己增加 3 点防御", "短臂剑 + 很小钢 buckler。"],
            ["S-40", "Felicity", "1", "基础", "攻击", "立即造成 3 点伤害，并为自己增加 4 点防御", "单手斧 + 小圆盾。"],
            ["S-41", "Riley", "2", "基础", "攻击", "立即造成 7 点伤害，并为自己增加 9 点防御", "手半剑 + 中等鸢盾。"],
            ["S-42", "Daniel", "3", "基础", "攻击", "立即造成 12 点伤害，并为自己增加 15 点防御", "重战锤 + 最大加热盾。半框眼镜。"],
        ],
        card_widths,
        card_center,
    )

    add_heading(doc, "采购成员卡 P-26～P-37", 1)
    add_body(doc, "重板甲盾兵立绘。不持攻击兵器。盾面按防御强弱缩放，每人不同纹章。同名销售卡是另一张牌。基础卡每局带 2 张，收藏卡需胜利解锁。")
    add_flex_table(
        doc,
        card_headers,
        [
            ["P-26", "Hermione", "0", "收藏", "防御", "为自己增加 3 点防御", "小加热盾：紫底金边，正中铃兰。"],
            ["P-27", "Kinyond", "1", "基础", "防御", "为自己增加 6 点防御", "中等鸢盾：黑底金鸢尾，尖端小锤记。"],
            ["P-28", "Harry", "2", "基础", "防御", "为自己增加 13 点防御", "最大加热盾：深蓝底大金太阳。半框眼镜。"],
            ["P-29", "Kade", "0", "基础", "防御", "为一名玩家增加 2 点防御", "小圆盾：白底蓝十字，交叉处伸出的手。"],
            ["P-30", "Sandra", "1", "基础", "防御", "为一名玩家增加 5 点防御", "中等圆盾：绿底向外伸的手/羽翼。"],
            ["P-31", "Charlene", "2", "基础", "防御", "为一名玩家增加 11 点防御", "大加热盾：酒红底银月+斜带。酒红兜帽。"],
            ["P-32", "Ethan", "0", "收藏", "防御", "为双方各增加 1 点防御", "极小黄铜 buckler：正中「1」+ 计数刻痕。"],
            ["P-33", "Carl", "1", "收藏", "防御", "为双方各增加 3 点防御", "targe：深棕皮面交错斜带。"],
            ["P-34", "Colin", "2", "收藏", "防御", "为双方各增加 7 点防御", "中等加热盾：左蓝右金对分，两边各一星。"],
            ["P-35", "Amy", "3", "基础", "防御", "为双方各增加 12 点防御", "大 pavise：蓝底双狮相对。"],
            ["P-36", "Duane", "0", "收藏", "防御", "替队友承受下一次打向该队友的霸凌者攻击", "立地高 pavise，城门涂鸦「挡住」。自己仍吃打向自己的那下；队友那一下改打到出牌者，护盾照常结算。"],
            ["P-37", "Kad3", "1", "收藏", "防御", "自己防御 -1，为队友增加 10 点防御", "大加热盾：左素钢右金漆，红线缝住裂缝。自盾不低于 0。"],
        ],
        card_widths,
        card_center,
    )

    add_heading(doc, "顾客效果", 1)
    add_body(doc, "开局按登场权重随机一名顾客。顾客效果每回合开始独立掷一次（第 1 回合不掷）。不会因此复活已倒下的护卫。")
    add_flex_table(
        doc,
        ["顾客", "编码", "登场权重", "效果触发", "顾客效果", "人话"],
        [
            ["善良稳重", "CUSTOMER_KIND", "32（约 29%）", "约 30%", "本回合霸凌者攻击 -1", "偶尔把这一刀削轻一点。"],
            ["胆小怕事", "CUSTOMER_TIMID", "38（约 35%）", "约 65%", "霸凌者血量 +2（上限同步提高）", "最常见。局会被拖长。"],
            ["焦虑难安", "CUSTOMER_ANXIOUS", "30（约 27%）", "约 60%", "本回合霸凌者攻击 +2", "这一拍刀更大。"],
            ["闲逛双客", "CUSTOMER_WINDOW", "10（约 9%）", "约 20%", "两名存活护卫各恢复 2 点血", "结伴来的，奶也是两人份，救不活倒下的人。"],
        ],
        [Cm(2.8), Cm(3.6), Cm(3.2), Cm(2.4), Cm(6.2), Cm(5.8)],
        {2, 3},
    )

    add_heading(doc, "霸凌者特效", 1)
    add_body(doc, "本局霸凌者锁死跟顾客走，不再另外随机。身份类常驻，加码类概率。旧模板「普通霸凌者」已停用。")
    add_flex_table(
        doc,
        ["霸凌者", "编码", "触发", "特效", "情报页"],
        [
            ["点名恶霸", "BULLY_FOCUS_001", "常驻", "每回合只打当前血更低的那人；持平打销售。攻击数值不另加。", "每回合专打更弱的护卫。"],
            ["硬扛恶霸", "BULLY_SHIELD_001", "约 40% / 回合", "本回合获得 4 点可打掉的盾；未触发则没有。下回合开始先清残留再重新掷。", "约四成回合胸口会多一层 4 点盾。"],
            ["针对恶霸", "BULLY_REVENGE_001", "约 50% / 回合", "仍打两人。若触发，本回合对霸凌者造成伤害（扣血）更高的那人挨打额外 +2；持平点销售。按本回合增量，不是整局累计。", "约一半回合会盯打得最疼的人加一刀。"],
            ["不落单恶霸", "BULLY_PAIR_001", "双人零掉血后再 35%", "仍打两人。若本回合两名存活者都没有掉血，再掷 35%：下回合各追加一次半伤（先打整刀，护盾吃完后再打半刀）。", "两人都挡住时，约三成会在下一拍再抽一次。"],
        ],
        [Cm(2.8), Cm(3.6), Cm(3.4), Cm(8.8), Cm(5.4)],
        {2},
    )

    add_heading(doc, "一客一霸凌者", 1)
    add_flex_table(
        doc,
        ["顾客", "顾客做什么", "绑定霸凌者", "霸凌者做什么", "为什么这样配"],
        [
            ["善良稳重", "约 30% 攻 -1", "点名恶霸", "每回合点名更弱的人", "顾客在减刀，威胁改成专挑软的。逼采购给血少的人罩，Duane 有事做。"],
            ["胆小怕事", "约 65% 霸凌者 +2 血", "硬扛恶霸", "约 40% 本回合 4 盾", "顾客已经在拖血量。偶发盾让销售要想这回合要不要把盾打穿。"],
            ["焦虑难安", "约 60% 攻 +2", "针对恶霸", "约 50% 盯输出 +2", "顾客已经加攻，不再全局加伤，改成谁还手打谁。"],
            ["闲逛双客", "约 20% 两人各回 2", "不落单恶霸", "全挡住后再 35% 下回合半伤", "双客结伴、回血也是两人份，霸凌者不肯放过任何一边。"],
        ],
        [Cm(2.6), Cm(3.6), Cm(2.8), Cm(5.2), Cm(9.8)],
        set(),
    )

    add_heading(doc, "结算补充", 2)
    for line in [
        "点名恶霸只打一人时，另一人本回合不挨打。若被点名的人有 Duane 守护，这一刀改打到 Duane。",
        "硬扛恶霸的 4 点盾可被攻击牌和延迟伤害打掉；未触发的回合开场盾为 0。",
        "针对恶霸比较的是本回合实际扣掉的霸凌者血量。打在盾上的伤害不计入。",
        "不落单恶霸要两名存活者本回合血量损失都为 0 才掷 35%。有人被 Duane 替伤而自己没掉血、另一人挨了打，不会触发。",
        "不要把顾客骰和霸凌者骰绑在一起。连续空骰或连续双中都可能发生。",
    ]:
        add_body(doc, "· " + line)

    add_heading(doc, "修订记录", 2)
    rec = doc.add_table(rows=3, cols=2)
    rec.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(rec.rows[0].cells[0], "日期", header=True, center=True)
    write_cell(rec.rows[0].cells[1], "说明", header=True, center=True)
    write_cell(rec.rows[1].cells[0], "2026-09-01", center=True)
    write_cell(rec.rows[1].cells[1], "新建。收录当前对局成员卡 S-35～S-43、P-26～P-37。")
    write_cell(rec.rows[2].cells[0], "2026-09-01", center=True)
    write_cell(rec.rows[2].cells[1], "写入四名顾客现有效果，以及一对一绑定的点名/硬扛/针对/不落单恶霸特效。")

    save_doc(doc, CURRENT_OUT)


if __name__ == "__main__":
    import sys
    if "v2" in sys.argv:
        main_v2()
    elif "current" in sys.argv:
        main_current()
    else:
        main()
